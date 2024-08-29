from docx import Document
import re

def print_document_with_placeholders(file_path):
    doc = Document(file_path)
    
    # Define a regex pattern for placeholders (e.g., {{placeholder}})
    placeholder_pattern = re.compile(r'\{\{.*?\}\}')

    print("Document Content:")
    
    # Check paragraphs for placeholders
    for para in doc.paragraphs:
        text = para.text
        if placeholder_pattern.search(text):
            print(f"Paragraph with placeholder: {text}")

    # Check tables for placeholders
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                if placeholder_pattern.search(text):
                    print(f"Table cell with placeholder: {text}")

print_document_with_placeholders('atska.docx')
