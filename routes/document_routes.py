from flask import Blueprint, render_template, request, redirect, url_for, send_file, flash
from io import BytesIO
from utils.db import get_db_connection
from docx import Document
from utils.docx_utils import replace_placeholders, extract_placeholders_and_data, extract_table_data, combine_data, insert_data_into_table, replace_placeholders_with_sums

document_routes = Blueprint('document', __name__)

@document_routes.route('/select-documents', methods=['GET'])
def select_documents():
    """Render the page to select documents."""
    with get_db_connection() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT id, name, date FROM documents")
        documents = cursor.fetchall()
    print(f"Documents fetched: {documents}")  # Debugging line
    return render_template('select_documents.html', documents=documents)

@document_routes.route('/generate-from-selected', methods=['POST'])
def generate_from_selected():
    """Process selected documents."""
    doc_ids = request.form.getlist('doc_ids')
    print(f"Received document IDs: {doc_ids}")

    if not doc_ids:
        flash("No documents selected.", "warning")
        return redirect(url_for('document.select_documents'))

    try:
        with get_db_connection() as conn:
            cursor = conn.cursor()
            placeholders = []
            for doc_id in doc_ids:
                cursor.execute("SELECT name FROM documents WHERE id = %s", (doc_id,))
                doc_name = cursor.fetchone()
                if doc_name:
                    first_part = doc_name[0].split('_')[0]  # Extract the first part before the underscore
                    placeholders.append(first_part)

        combined_name = "_".join(placeholders)  # Combine all first parts with underscores
        output_path = process_documents(doc_ids, combined_name)
        flash("Documents processed successfully.", "success")
        return send_file(output_path, as_attachment=True)
    except Exception as e:
        flash(f"An error occurred: {str(e)}", "danger")
        return redirect(url_for('document.select_documents'))

def process_documents(doc_ids, combined_name):
    print(f"Processing documents with IDs: {doc_ids}")

    template_path = 'docs/Template.docx'

    with get_db_connection() as conn:
        cursor = conn.cursor()

        all_placeholders = {}
        all_replacements = {}

        # Collect all placeholders and replacements from the selected documents
        for doc_id in doc_ids:
            cursor.execute("SELECT data FROM documents WHERE id = %s", (doc_id,))
            doc_data = cursor.fetchone()
            
            if not doc_data:
                print(f"No data found for document ID: {doc_id}")
                continue
            
            doc = Document(BytesIO(doc_data[0]))
            placeholders = extract_placeholders_and_data(doc)
            replaced_data = replace_placeholders(doc, all_replacements)

            for placeholder, texts in placeholders.items():
                if placeholder not in all_placeholders:
                    all_placeholders[placeholder] = texts
                else:
                    all_placeholders[placeholder].extend(texts)

            for placeholder, value in replaced_data.items():
                all_replacements[placeholder] = [value]

        try:
            print(f"Loading template from: {template_path}")
            template_doc = Document(template_path)
        except Exception as e:
            print(f"Error loading template: {e}")
            raise

        # Process tables in the template
        for table_idx in range(len(template_doc.tables)):
            combined_data = []
            for doc_id in doc_ids:
                cursor.execute("SELECT data FROM documents WHERE id = %s", (doc_id,))
                doc_data = cursor.fetchone()
                
                if not doc_data:
                    print(f"No data found for document ID: {doc_id}")
                    continue

                doc = Document(BytesIO(doc_data[0]))

                if table_idx < len(template_doc.tables):
                    data_table = doc.tables[table_idx]
                    data = extract_table_data(data_table, table_idx)
                    combined_data = combine_data(combined_data, data)

            insert_data_into_table(template_doc, table_idx, combined_data)

        # Construct the output filename
        output_filename = f"{combined_name}_KatedrosAtaskaita.docx"
        output_path = f"{output_filename}"

        try:
            print(f"Saving updated template to: {output_path}")
            template_doc.save(output_path)
        except Exception as e:
            print(f"Error saving updated template: {e}")
            raise

    print(f"Generated document saved at: {output_path}")
    return output_path



@document_routes.route('/replace-file/<int:doc_id>', methods=['POST'])
def replace_file(doc_id):
    """Replace a file in the database."""
    if 'file' not in request.files:
        return "No file part", 400

    file = request.files['file']
    if file.filename == '':
        return "No selected file", 400

    if file and file.filename.endswith('.docx'):
        file_data = file.read()
        
        with get_db_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("UPDATE documents SET data = %s WHERE id = %s", (file_data, doc_id))
            conn.commit()
        
        return "File replaced successfully", 200

    return "Invalid file type", 400

@document_routes.route('/form', methods=['GET'])
def form():
    """Render the form page."""
    return render_template('form.html')

@document_routes.route('/download-file/<int:doc_id>', methods=['GET'])
def download_file(doc_id):
    """Download a file from the database."""
    with get_db_connection() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT data, name FROM documents WHERE id = %s", (doc_id,))
        doc_data = cursor.fetchone()

    if doc_data:
        data, name = doc_data
        return send_file(BytesIO(data), as_attachment=True, download_name=name)
    return "Document not found.", 404
