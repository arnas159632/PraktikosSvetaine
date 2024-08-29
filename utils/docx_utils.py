from docx import Document

def replace_placeholders(doc, replacements):
    replaced_data = {}

    def replace_in_text(text, replacement):
        for placeholder, replacement_texts in replacement.items():
            replacement_text = ' '.join(replacement_texts)
            if placeholder in text:
                print(f"Replacing '{placeholder}' with '{replacement_text}'")
                text = text.replace(placeholder, replacement_text)
                replaced_data[placeholder] = replacement_text
        return text

    for paragraph in doc.paragraphs:
        original_text = paragraph.text
        paragraph.text = replace_in_text(paragraph.text, replacements)
        if paragraph.text != original_text:
            print(f"Updated paragraph text: {paragraph.text}")

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                original_text = cell.text
                new_text = ''
                for paragraph in cell.paragraphs:
                    new_text += replace_in_text(paragraph.text, replacements) + '\n'
                cell.text = new_text.strip()
                if cell.text != original_text:
                    print(f"Updated cell text: {cell.text}")

    return replaced_data

def extract_placeholders_and_data(doc):
    placeholders = {}

    def extract_from_text(text):
        parts = text.split()
        for part in parts:
            if part.startswith('{{') and part.endswith('}}'):
                placeholder = part
                if placeholder not in placeholders:
                    placeholders[placeholder] = []
                placeholders[placeholder].append(text)
                print(f"Found placeholder: {placeholder}")

    for paragraph in doc.paragraphs:
        extract_from_text(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                extract_from_text(cell.text)

    return placeholders

def extract_table_data(table, table_idx):
    table_data = []
    
    # Special handling for table index 16
    if table_idx == 16:
        rows_to_extract = [11, 12, 13]
        for row_idx in rows_to_extract:
            if row_idx < len(table.rows):
                row_data = [cell.text.strip() for cell in table.rows[row_idx].cells]
                table_data.append(row_data)
        print(f"Extracted specific rows {rows_to_extract} for table index {table_idx}: {table_data}")
    else:
        # Extract all rows except header if not table 16
        for row in table.rows[1:]:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        print(f"Extracted table data for table index {table_idx}: {table_data}")
    
    return table_data

def combine_data(data1, data2):
    combined_data = data1 + data2
    print(f"Combined data: {combined_data}")
    return combined_data

def insert_data_into_table(doc, table_idx, data):
    if table_idx in {14, 26,27, 28, 35,40, 47}:
        print(f"Table Index: {table_idx} is excluded from modifications.")
        return

    table = doc.tables[table_idx]

    if table_idx == 16:
        fixed_header_rows_count = 10
    else:
        headers = TABLE_HEADERS.get(table_idx, [])
        fixed_header_rows_count = find_header_rows(table, headers)

    total_rows_needed = fixed_header_rows_count + len(data)
    ensure_table_has_rows(table, total_rows_needed)

    clear_data_rows(table, fixed_header_rows_count)

    for i, row_data in enumerate(data):
        row_idx = fixed_header_rows_count + i
        if row_idx >= len(table.rows):
            print(f"Warning: Not enough rows in table {table_idx}. Adding row {row_idx + 1}.")
            table.add_row()

        row_cells = table.rows[row_idx].cells
        for j, cell_data in enumerate(row_data):
            if j < len(row_cells):
                row_cells[j].text = cell_data

def ensure_table_has_rows(table, num_rows):
    current_rows = len(table.rows)
    if current_rows < num_rows:
        for _ in range(num_rows - current_rows):
            table.add_row()
        print(f"Ensured table has {num_rows} rows.")

def clear_data_rows(table, start_row):
    for row in table.rows[start_row:]:
        for cell in row.cells:
            cell.text = ''

def find_header_rows(table, headers):
    header_row_count = 0
    for row in table.rows:
        row_data = [cell.text.strip() for cell in row.cells]
        if row_data in headers:
            break
        header_row_count += 1
    print(f"Found {header_row_count} header rows.")
    return header_row_count

def calculate_sum_from_table(doc, table_idx, value_column_idx, condition_column_idx, condition_value):
    table = doc.tables[table_idx]
    total_sum = 0
    for row in table.rows[1:]:
        cells = row.cells
        if len(cells) > condition_column_idx and len(cells) > value_column_idx:
            if cells[condition_column_idx].text.strip() == condition_value:
                try:
                    value = float(cells[value_column_idx].text.strip())
                    total_sum += value
                except ValueError:
                    pass
    print(f"Calculated sum for table {table_idx}: {total_sum}")
    return total_sum

def replace_placeholders_with_sums(doc, placeholders):
    replaced_data = {}
    for placeholder in placeholders:
        if placeholder.startswith('{{sum') and placeholder.endswith('}}'):
            try:
                table_idx = int(placeholder[4:-2])  # For example, '{{sum11}}' -> table_idx = 11
                sum_result = calculate_sum_from_table(doc, table_idx, 4, 1, "Mokymo(si) kompetencijos")
                replaced_data[placeholder] = str(sum_result)
                print(f"Replaced placeholder '{placeholder}' with sum '{sum_result}'")
            except ValueError:
                print(f"Invalid placeholder format: {placeholder}")

    for paragraph in doc.paragraphs:
        for placeholder, value in replaced_data.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)
                print(f"Updated paragraph text with sum: {paragraph.text}")

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for placeholder, value in replaced_data.items():
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, value)
                        print(f"Updated cell text with sum: {cell.text}")

    return replaced_data

def print_table_data(doc):
    max_tables_to_print = 5
    for table_idx, table in enumerate(doc.tables[:max_tables_to_print]):
        if table_idx == 13:
            print(f"Table Index: {table_idx} is excluded from printing.")
            continue

        print(f"\nTable Index: {table_idx}")
        
        # Determine if the table has a fixed header
        if table_idx == 16:
            header_row_count = 10
        else:
            header_row_count = 0
            headers = TABLE_HEADERS.get(table_idx, [])
            if headers:
                header_row_count = find_header_rows(table, headers)

        # Print table rows, skipping header rows if necessary
        for row_idx, row in enumerate(table.rows):
            row_data = [cell.text.strip() for cell in row.cells]
            if table_idx != 16 or row_idx >= header_row_count:
                if row_idx == 0 or row_data != headers:
                    print(f"Row Index: {row_idx}, Row Data: {row_data}")

TABLE_HEADERS = {
    # Example:
    # 1: [['Header1', 'Header2'], ['Subheader1', 'Subheader2']],
    # 3: [['HeaderA', 'HeaderB']],
}

if __name__ == "__main__":
    doc_path = 'sample.docx'
    doc = Document(doc_path)
    print(f"Loaded document from: {doc_path}")

    replacements = {
        # Define any other replacements if needed
    }

    placeholders = extract_placeholders_and_data(doc)
    print(f"Extracted Placeholders: {placeholders}")

    replace_placeholders_with_sums(doc, placeholders)
    replace_placeholders(doc, replacements)
    
    print_table_data(doc)
    
    updated_doc_path = 'updated_sample.docx'
    doc.save(updated_doc_path)
    print(f"Saved updated document to: {updated_doc_path}")
